"""
代码中有从网络上参考或者直接复制过来的内容，对原作者致以谢意！
但因时间有些久了，来源无法一一追溯，故没有列出来源。报以歉意！
"""
# simtext相似度：
# simtext可以计算两文档间四大文本相似性指标，分别为：
#     Sim_Cosine cosine相似性
#     Sim_Jaccard Jaccard相似性
#     Sim_MinEdit 最小编辑距离
#     Sim_Simple 微软Word中的track changes
import sys
import re
import string
from zhon.hanzi import punctuation as chinese_punctuation  # 中文标点符号
from docx.shared import RGBColor
from docx import Document
from docx.shared import Inches
import os
from simtext import similarity
from PyQt5.QtWidgets import (QWidget, QMainWindow, QTextEdit, 
    QAction, QFileDialog, QApplication, QPushButton,QCheckBox,
    QProgressBar,QGridLayout,QLabel,QLineEdit,QToolTip,
    QMessageBox,QPlainTextEdit)
import PyQt5.QtWidgets
from PyQt5.QtGui import QIcon,QFont,QPixmap  
from PyQt5.QtCore import Qt, pyqtSignal, QObject,QBasicTimer
from PyQt5 import QtCore
###需要后边的import，是因为打包要一起打包进去
import simtext
import zhon
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.metrics import jaccard_similarity_score
from difflib import Differ, SequenceMatcher
import jieba
from math import *
import warnings

findir  = 'documents/'      ### The path of the PDFs
foutdir = 'extracted_text/' ###


english_punctuation = string.punctuation  # 英文标点符号
chi_punc = '|'.join([c for c in chinese_punctuation])
eng_punc = '|'.join([c for c in english_punctuation])
punc = chi_punc + eng_punc
# 注意 punc 中的'||'会导致逐字符分句的情况，所以手动抛去；如果真要把‘|’也当做分隔符，再做研究
punc = punc[:-6]+punc[-4:]
#punc = punc.decode("utf-8")

def not_break(sen):
    return (sen != '\n' and sen != '\u3000' and  sen != '' and not sen.isspace())
def filter_data(ini_data):
    # ini_data是由句子组成的string
    new_data = list(filter(not_break, [data.strip() for data in ini_data]))
    return new_data
def split_sentence(para):
    my_stringList = filter_data(re.split(r''+("["+punc+"]"), para))
    """
    flags = [u"，",u"。",u"；",u"：",u"！",u"？",u'“',u"’"]
    for flag in flags:
         para2 = para.replace(flag,"\n")
    print(para2)
    para3=para2.split('\n')
    """
    return my_stringList 
    
 


def run_comparing(para1, para2, mymethod= "Sim_Cosine",mythreshold = 0.5):
    para1 = para1.replace('\n','')
    para2 = para2.replace('\n','')
    para12=para1   ### used for colored text
    para22=para2   ### used for colored text
    myoutput = open('检测结果.txt','w')
    
    fdoc = '检测结果.docx'
        
    # 创建文档对象
    # 打包时使用：
    document = Document()
    # 测试时使用(因为会出问题，所以要这样)：
    #document = Document( docx=os.path.join(os.getcwd(),'default.docx') )
    # 添加一级标题
    document.add_heading(u"使用的方法：%s, 阈值=%10.3f%%"%(mymethod,mythreshold*100.0),level = 1)
    
    """
    p=document.add_paragraph('')
    xxxx = "使用的方法：%s," % (mymethod)
    yyyy = "阈值=%10.3f%%"  %(mythreshold*100.0)
    run=p.add_run( xxxx+yyyy )
    run.font.color.rgb = RGBColor(250,0,0) 
    """
    
    document.add_paragraph('比较的两段为：')
    document.add_paragraph(para1, style = 'IntenseQuote')
    document.add_paragraph(para2, style = 'IntenseQuote')
    
    """
    # 添加表格: 1行3列
    table = document.add_table(rows = 1,cols = 2)
    # 获取第一行的单元格列表对象
    hdr_cells = table.rows[0].cells
    # 为每一个单元格赋值
    # 注：值都要为字符串类型
    hdr_cells[0].text = "句子"
    hdr_cells[1].text = "相似度"
    """
                
                
    mysentences = split_sentence(para1)
    print('要查询的文本中句子数目 =',len(mysentences))
    textDB      = split_sentence( para2.replace('\n','') )
    
    print("使用的方法：", mymethod, ', 阈值=', mythreshold*100.0 , '%',file=myoutput)
    print("段落1：",file=myoutput)
    print(para1,file=myoutput)
    print("段落2：",file=myoutput)
    print(para2,file=myoutput )
    
    text1color=""
    text2color=""
    it2 = 0
    for text2x in textDB:
        it2+=1
        text2color=text2color+  '<font color="red">[' + str(it2) + ']</font>' +text2x
    
    mean_p_1=0.0
    i1=0
    pmarker=0.0
    for sentence in mysentences:
        pmarker=0.0
        sim = similarity()
        it2=0
        imatch=0
        imatch2=0;  textimatch2=''
        i=0
        for text in textDB:
            it2+=1
            res = sim.compute(text, sentence)
            
            i+=1
            
            if pmarker<res[mymethod]:
                pmarker=res[mymethod]
                imatch=it2
                imatch2=i
                textimatch2=text
                
            if res[mymethod]>=mythreshold :
                mean_p_1 += res[mymethod]*100.0
                if pmarker<res[mymethod]:
                    pmarker=res[mymethod]
                    
                i1+=1
                print("#############################################################", file=myoutput)
                print('相似度    =%10.5f %%'%(res[mymethod]*100.0), file=myoutput )
                print("句子1：",sentence, file=myoutput)
                print('句子2：  ',text, file=myoutput)
                
                
                # 往文档中添加段落
                p = document.add_paragraph('相似度    =')
                run=p.add_run('%10.5f %%'%(res[mymethod]*100.0))
                run.font.color.rgb = RGBColor(250,0,0)
                p.add_run('， 两个句子如下：').bold = True
                #p.add_run('and some ')
                #p.add_run('italic.').italic = True
                document.add_paragraph(sentence, style =  'ListBullet')#'ListNumber')#style = 'IntenseQuote')
                document.add_paragraph(text, style =  'ListBullet')#'ListNumber')#style = 'IntenseQuote')
                
                
                """
                # 为表格添加一行
                new_cells = table.add_row().cells
                new_cells[0].text = sentence
                new_cells[1].text = ''
                new_cells = table.add_row().cells
                new_cells[0].text = text
                new_cells[1].text = '%10.5f %%'%(res[mymethod]*100.0)
                """
                
        pc = "(%5.2f%%,%d)"%(pmarker*100.0,imatch)

        #print(pmarker)    
        if pmarker>=mythreshold:
            text1color = text1color + '<font color="red">'+ sentence+pc+'</font> <font color="cyan">|</font> '
            xxxx = '<font color="red">'+ sentence+'</font>'
            #print(sentence,xxxx)
            para12 = para12.replace(sentence,xxxx)
            yyyy = '<font color="red">'+ textimatch2 +'</font>'
            para22 = para22.replace(textimatch2,yyyy)
            #print(para12,sentence,xxxx)
        else:
            text1color = text1color + '<font color="blue">'+ sentence + pc + '</font> <font color="cyan">|</font> '
            xxxxx = '<font color="black">'+ sentence+'</font>'
            para12 = para12.replace(sentence,xxxxx)
            #yyyyy = '<font color="black">'+ textimatch2+'</font>'
            #para22 = para22.replace(textimatch2,yyyyy)
            
            
    text1color = text1color + '\n\n  <font color="black">说明：括号内（第一项为相似度，第二项为与之相似的句子编号)。详情请查看 “检测结果.txt“ 或者 “检测结果.docx“。</font> <font color="red">可能不准确！仅供参考！</font>' 
    
    try:
        mean_only = mean_p_1/float(i1)
    except:
        mean_only = 0.0
        
                
    print('**************************', file=myoutput) 
    print('大于阈值的文本的平均相似度= %10.5f %%'%(mean_only) , file=myoutput)           
    #print('整体平均相似度= %10.5f %%'%(mean_all) , file=myoutput) 
    #print("注意整体相似度为所有句子的平均相似度（包含相似度小于阈值的内容）", file=myoutput)
    print('**************************', file=myoutput)      
    
    """
    col_width = [6.5,0.5]
    for col in range(2):
        table.cell(0,col).width = Inches(col_width[col])
    """ 
     
    p = document.add_paragraph('')
    run= p.add_run('大于阈值的文本的平均相似度= %10.5f %%'%(mean_only) ) 
    run.font.color.rgb = RGBColor(250,0,0)    
    #p.add_run("\n注意整体相似度为所有句子的平均相似度（包含相似度小于阈值的内容）").bold = True
    
    """
    p = document.add_paragraph('')
    run= p.add_run('整体平均相似度= %10.5f %%'%(mean_all) ) 
    run.font.color.rgb = RGBColor(250,0,0)    
    p.add_run("\n注意整体相似度为所有句子的平均相似度（包含相似度小于阈值的内容）").bold = True
    """
    
    document.save(fdoc)
    
    myoutput.close()
    
    return mean_only, text1color, text2color, para12, para22


class ComparetwoParagraphGUI(QMainWindow):

    def __init__(self):
        super().__init__()

        self.initUI()


    def initUI(self):   
        
        QToolTip.setFont(QFont('arial', 10))
        self.setToolTip('两段文字的相似度对比（作者：星未 [知乎]）') 
        
        ### set common font:
        #self.setFont(QFont('SansSerif', 10))
        
        
        self.statusBar()
        key_bg_color='#ff6600'
        
        label_bg_color="#ffc266"#"#ffe0cc" #'#e6e6e6'
        
        ###########################
        window = self.setGeometry(300, 300, 350, 600)
        self.setWindowTitle('两段文字的相似度对比（星未）')
        self.setWindowIcon( QIcon('icon.ico') )
        
        ##########################

        self.text1 = QTextEdit(self)
        self.text1.setGeometry(10, 0, 330, 150) #(x,y,xwidth,ywidth)
        self.text1.setText("")
        self.text1.setPlaceholderText('输入自己写的文字。这里的文字将逐句与下面的文字比较。\n相似度高于阈值的句子将被标红！')
        self.text1.setToolTip("")
        #self.text1.setWordWrap(True)
        #self.text1.setEnabled(False)
        
        
        ##########################

        self.text2 = QTextEdit(self)
        self.text2.setGeometry(10, 150, 330, 150)
        #self.save_tip = self.save_text 
        self.text2.setText("")
        self.text2.setPlaceholderText('输入要对比的文字。')
        self.text2.setToolTip( "" )
        #self.text2.setWordWrap(True)
        
        ##########################

        self.text3 = QTextEdit(self)
        self.text3.setGeometry(10, 300, 330, 150)
        self.text3.setText("")
        self.text3.setPlaceholderText('此处显示结果。')
        self.text3.setToolTip( "" )
        #self.text3.setEnabled(False)
        
        
        ###########################
        self.QL_threshold = QLabel(self)
        self.QL_threshold.setGeometry(10, 460, 150, 30)
        self.QL_threshold.setAutoFillBackground(True)
        self.QL_threshold.setText('相似度阈值（%）：') 
        
        ###########################
        self.threshold = QTextEdit(self)
        self.threshold.setGeometry(160, 460, 150, 30)
        self.threshold.setText("50")
        self.threshold.setToolTip( "输入阈值。注意：相似度大于此阈值的句子将被保存入检测结果文件！" )
        
        ### output information, such as Running, CPUtime. 
        self.output_info = QLabel(self)
        self.output_info.setGeometry(10, 500, 200, 90)
        self.output_info.setAutoFillBackground(True)
        self.output_info.setText('粘贴两段文字后，点击“对比”')
        self.output_info.setStyleSheet("background-color: "+key_bg_color+";  border: 1px solid "+key_bg_color) 
        self.output_info.setToolTip('该区域显示进程，结果仅供参考！')
        self.output_info.setWordWrap(True)  
        
        ###########################
        self.btn_start = QPushButton('对比', self)
        self.btn_start.setGeometry(QtCore.QRect(220, 500, 120, 90))
        self.btn_start.clicked.connect(self.change_output_info_to_running)
        self.btn_start.clicked.connect(self.runModel)
        self.btn_start.setToolTip('点击开始对比')
        self.btn_start.setStyleSheet("background-color: cyan;  border: 1px solid "+key_bg_color) 
        
        
        #window.setToolTip('A two-phase gas-grain chemical model.')
        #self.setStyleSheet("background-color: gray;  border: 1px solid gray") 
        
        self.show()
    def change_output_info_to_NO_text(self,xtext):
        self.output_info.setText('文本框 %s 中，文本不能为空！'% (xtext) )
        QApplication.processEvents() 
        
    def change_output_info_to_running(self):
        self.output_info.setText('对比中，请等待。。。')
        QApplication.processEvents() 
        
    def change_output_info_to_results(self,mean_only):
        self.output_info.setText('大于阈值的平均相似度=%5.2f%%'%(mean_only))
        QApplication.processEvents() 

    def runModel(self):
        #mytext      = '北极地区冻结主权是不可行的，北极治理构变动，丹麦虽然并不是世界上最靠近北极圈的国家，但它也正试图将北极纳入自己的版图。俄罗斯联邦国家杜马副主席奇林加罗夫率领北极海洋科考团,我吃北极海洋考察'
        #para        = "北极地区冻结主权不可行的，北极治理的智缘结构变动，丹麦虽然并不是世界北极圈的国家，但它也正试北极纳的版图。俄罗斯联邦国主席奇林领北极海洋科考团"

        mytext = self.text1.toPlainText()
        para   = self.text2.toPlainText()
        #print(mytext)
        threshold = float(self.threshold.toPlainText())/100.0
        
        try:
            if mytext=="" and para!="":
                self.text1.setPlaceholderText('文本不能为空！')
                self.change_output_info_to_NO_text("1")
                self.text3.setPlaceholderText('输入文本不能为空！')
            elif para=="" and mytext!="":
                self.text2.setPlaceholderText('文本不能为空！')
                self.change_output_info_to_NO_text("2")
                self.text3.setPlaceholderText('输入文本不能为空！')
            elif para=="" and mytext=="":
                self.text2.setPlaceholderText('文本不能为空！')
                self.text1.setPlaceholderText('文本不能为空！')
                self.change_output_info_to_NO_text(" 1 和 2 ")
                self.text3.setPlaceholderText('输入文本不能为空！')
            elif mytext!="" and para!="":
                mean_only, text1color, text2color, para1mk, para2mk = run_comparing(mytext, para, mythreshold=threshold)
                self.change_output_info_to_results(mean_only)
                
                self.text1.setText('')
                self.text1.setText(para1mk)
                
                self.text2.setText('')
                self.text2.setText(para2mk)
            
                #print(para12)
                
                self.text3.setText("")
                self.text3.setText(text1color)#+'比较文的标号：'+text2color)
        except:
            pass
        ##self.output_info.setText('Running')
        #msg = QMessageBox()
        #msg.setIcon(QMessageBox.Information)
        #msg.setText("对比完成，请查看 “检测结果.txt” 或者 “检测结果.docx” ")
        #msg.exec_()
        

app = QApplication(sys.argv)
app.setWindowIcon(QIcon('icon.ico'))
ex  = ComparetwoParagraphGUI()
sys.exit(app.exec_())

#input("please input any key to exit!")  ###防止exe闪退
